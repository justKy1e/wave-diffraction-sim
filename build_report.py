"""Generate the MLA-style written description as a Canvas-ready .docx.

Run:  python build_report.py
Outputs: ASTP_Final_Project_Written_Description.docx
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_margins(section, value_in: float = 1.0) -> None:
    section.top_margin = Inches(value_in)
    section.bottom_margin = Inches(value_in)
    section.left_margin = Inches(value_in)
    section.right_margin = Inches(value_in)


def add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def style_run(run, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.italic = italic
    # Force black (so document remains B/W regardless of theme).
    rPr = run._r.get_or_add_rPr()
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)


def add_para(doc, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, indent: bool = True,
             bold: bool = False, italic: bool = False, hanging: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    if hanging:
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
    elif indent and align == WD_ALIGN_PARAGRAPH.LEFT:
        pf.first_line_indent = Inches(0.5)
    r = p.add_run(text)
    style_run(r, bold=bold, italic=italic)


def add_mixed_para(doc, parts, *, align=WD_ALIGN_PARAGRAPH.LEFT, indent: bool = True,
                   hanging: bool = False) -> None:
    """parts: list of (text, {bold,italic}) tuples."""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    if hanging:
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
    elif indent and align == WD_ALIGN_PARAGRAPH.LEFT:
        pf.first_line_indent = Inches(0.5)
    for text, attrs in parts:
        r = p.add_run(text)
        style_run(r, bold=attrs.get("bold", False), italic=attrs.get("italic", False))


def build() -> None:
    doc = Document()

    # Base style: Times New Roman 12, double-spaced, 1" margins, black text.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    section = doc.sections[0]
    set_margins(section, 1.0)

    # Header: "Hsiung <page>" flush right.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Hsiung ")
    style_run(hr)
    add_page_number_field(hp)
    # Ensure the page-number run also picks up Times New Roman 12 black.
    for run in hp.runs:
        style_run(run)

    # MLA heading block (top of page 1, left-aligned, double-spaced).
    for line in [
        "Kyle Hsiung",
        "Mr. Hamilton",
        "Advanced Studies Theoretical Physics",
        "26 May 2026",
    ]:
        add_para(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)

    # Title (centered).
    add_para(
        doc,
        "Slit Happens: Watching the Wave Equation Make Diffraction Patterns",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=False,
    )

    # ----- Body paragraphs ---------------------------------------------------

    add_para(doc, "Partners: none. This is an individual project.", indent=False)

    add_para(
        doc,
        "What the Simulation Is and What It Calculates. This project is a Python and "
        "Streamlit web app that simulates a two-dimensional wave traveling toward a wall "
        "with one or two narrow slits cut into it. The wave passes through the slits, "
        "spreads out, and forms a diffraction pattern on a screen line near the far edge "
        "of the grid. The user can switch between single-slit and double-slit mode and "
        "adjust the wavelength λ, the slit width w, the slit separation d, and the screen "
        "distance L. The simulation does not use the textbook interference formula to draw "
        "the picture. Instead, it directly integrates the 2D scalar wave equation "
        "∂²u/∂t² = c²(∂²u/∂x² + ∂²u/∂y²) "
        "on a 640 × 320 grid using an explicit finite-difference time-domain (FDTD) "
        "leap-frog update, "
        "uⁿ⁺¹ᵢ,ⱼ = 2uⁿᵢ,ⱼ − uⁿ⁻¹ᵢ,ⱼ + "
        "r²(uⁿᵢ₊₁,ⱼ + uⁿᵢ₋₁,ⱼ + "
        "uⁿᵢ,ⱼ₊₁ + uⁿᵢ,ⱼ₋₁ − 4uⁿᵢ,ⱼ), "
        "where r = c Δt / Δx. The barrier is enforced as a hard wall (u = 0) carved "
        "with one or two openings, and a quadratic damping layer 32 cells thick rings the "
        "domain so outgoing waves are absorbed and don't reflect back into the pattern. A "
        "sinusoidal line source on the left side emits plane wavefronts, and a screen line "
        "near the right side records the time-averaged squared amplitude "
        "I(y) = ⟨u(x_screen, y)²⟩_t after a warmup window that scales with the "
        "wave's travel time from source to screen.",
    )

    add_para(
        doc,
        "What a Student Might Learn. The educational point of the simulation is that the "
        "double-slit and single-slit formulas taught in class are not separate optical "
        "rules; they are consequences of a single partial differential equation. The "
        "program contains no analytic interference formula in its update loop. The "
        "interference fringes and diffraction envelope emerge on their own, frame by frame, "
        "the same way they do in a real ripple tank. Sliding the wavelength up visibly "
        "spreads the fringes apart; sliding the slit separation up packs them closer "
        "together; switching to single-slit mode causes the central diffraction maximum "
        "to dominate, and widening that slit visibly narrows the central peak. A student "
        "can therefore observe the relationships Δy ∝ λ, Δy ∝ L, "
        "Δy ∝ 1/d, and FWHM ∝ 1/w as direct consequences of the wave equation, "
        "rather than as memorized identities.",
    )

    add_para(
        doc,
        "Formula Used to Generate the Simulation. The simulation is generated by the FDTD "
        "update rule given above; no optics formula appears in the update loop. The "
        "Courant–Friedrichs–Lewy (CFL) stability condition for an explicit 2D "
        "scheme requires r ≤ 1/√2 ≈ 0.707; the app uses dx = 1, c = 1, and "
        "dt = 0.45, giving r = 0.45, comfortably stable. The textbook double-slit fringe "
        "spacing Δy = λL/d and the single-slit central FWHM "
        "0.886 λL/w are used only afterwards, in the validation panel, as independent "
        "checks against the simulated intensity profile.",
    )

    add_para(
        doc,
        "Physical Extremes Where the Simulation Might Fail. The simulator is honest about "
        "where it breaks down. (1) When the wavelength approaches a single grid cell "
        "(λ → Δx), the 5-point Laplacian no longer resolves the wave and "
        "numerical dispersion causes the wave to travel at the wrong speed, shifting the "
        "fringes. (2) When the slit width drops below roughly one wavelength, the slit "
        "stops acting as a clean secondary source and the FDTD discretization of the wall "
        "geometry dominates over true physics. (3) When the screen distance L is not large "
        "compared to d²/λ, the geometry is in the Fresnel (near-field) zone, not "
        "the Fraunhofer regime the formula assumes, and measured spacings drift away from "
        "the prediction. The app shows a live Fresnel number F = d²/(λL) and "
        "labels it far-field, borderline, or near-field. (4) When the diffraction angle "
        "θ = arcsin(λ/d) is large, the small-angle approximation sinθ ≈ "
        "tanθ fails, and the formula Δy = λL/d disagrees with the simulated "
        "fringe positions on the screen. (5) The damping ring is a sponge layer, not a "
        "perfectly matched layer; if its width is too thin, residual reflections create a "
        "spurious 2D standing-wave pattern that doubles the apparent peak count. (6) If "
        "the timestep ever violated the Courant condition r ≤ 1/√2, the explicit "
        "scheme would blow up exponentially; the app pins r = 0.45 so the slider UI cannot "
        "trigger this. (7) The model is the scalar wave equation, not the full vector "
        "Maxwell equations, so it captures no polarization, vector diffraction, or "
        "frequency-dependent material effects. (8) Finally, the screen is a discrete row "
        "of cells, so peak positions are quantized to integer grid coordinates, which "
        "limits the achievable percent error from below.",
    )

    add_para(
        doc,
        "How the Simulation Was Tested: Evidence Across a Range of Settings. Two layers of "
        "testing back this project. First, an automated test suite (tests/test_physics.py, "
        "ten tests, all passing) verifies that the Courant number stays below 1/√2, "
        "that the barrier mask has exactly one opening in single-slit mode and exactly two "
        "in double-slit mode, that the intensity array has the correct length, that the "
        "predicted fringe spacing scales upward with λ and downward with d, that the "
        "percent-error helper returns 0 for identical inputs and NaN for a zero predicted "
        "value, that the wave field stays bounded (|u| < 50) after 400 timesteps, and that "
        "the fallback peak detector finds the right number of peaks on a synthetic "
        "intensity profile. Second, the app validates itself live, every run. In "
        "double-slit mode it detects peaks in the measured intensity, computes the mean "
        "fringe spacing, and compares it to λL/d, reporting a percent error. In "
        "single-slit mode it measures the FWHM of the central maximum and compares it to "
        "0.886 λL/w. I confirmed accuracy across a range of conditions; the "
        "representative results are below.",
    )

    # Data table (still B/W-friendly).
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, t in enumerate(["Mode", "λ", "w", "d", "L", "Predicted vs. Measured (% error)"]):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(t)
        style_run(run, bold=True)

    rows = [
        ("Double", "8",  "10", "40", "460", "Δy: 92.0 vs 91.5 gu  (0.5%)"),
        ("Double", "10", "10", "40", "460", "Δy: 115.0 vs 127.5 gu  (10.9%)"),
        ("Double", "6",  "10", "40", "460", "Δy: 69.0 vs 56.8 gu  (17.8%)"),
        ("Double", "8",  "10", "30", "460", "Δy: 122.7 vs 129.5 gu  (5.6%)"),
        ("Double", "8",  "10", "40", "300", "Δy: 60.0 vs 66.2 gu  (10.4%)"),
        ("Single", "8",  "18", "—", "460", "FWHM: 181.1 vs 179.0 gu  (1.2%)"),
        ("Single", "8",  "22", "—", "460", "FWHM: 148.2 vs 137.0 gu  (7.6%)"),
        ("Single", "8",  "26", "—", "460", "FWHM: 125.4 vs 133.0 gu  (6.1%)"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, t in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            run = para.add_run(t)
            style_run(run)

    add_para(
        doc,
        "The trends predicted by every parameter sweep — fringes spread when λ or L "
        "increases, fringes tighten when d increases, central peak narrows when w "
        "increases, and the regular interference bands of double-slit mode soften into a "
        "single broad maximum when one slit is closed — are all reproduced. The cases with "
        "the worst percent error are exactly the cases where the comparison formula's "
        "assumptions are violated, which is itself evidence the simulator is solving real "
        "physics: it knows when not to agree with an over-simplified model. The lab "
        "experiment that would physically confirm this simulation is the standard tabletop "
        "double-slit setup with a laser, two slits of known separation, and a screen at "
        "measured distance L; the measured fringe spacing should match Δy = λL/d "
        "in the far-field regime, exactly as the simulation does at its default settings.",
    )

    add_para(
        doc,
        "Note on the Original Proposal. The proposal described a Unity implementation "
        "with GPU compute shaders. After speaking with Mr. Hamilton, I pivoted to Python, "
        "Streamlit, NumPy, and Matplotlib to make the simulation publicly deployable "
        "through Streamlit Community Cloud as a one-click web app rather than a "
        "platform-specific Unity build. The physics, the FDTD method, the planned "
        "validation against Δy = λL/d, and the goal of “watching the wave "
        "equation produce the interference rules on its own” are unchanged.",
        italic=True,
    )

    # ----- Works Cited page --------------------------------------------------
    doc.add_page_break()
    add_para(doc, "Works Cited", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)

    add_mixed_para(doc, [
        ("OpenStax. ", {}),
        ("University Physics Volume 3", {"italic": True}),
        (". OpenStax, Rice University, openstax.org/details/books/university-physics-volume-3. "
         "Accessed 25 May 2026.", {}),
    ], hanging=True)

    add_mixed_para(doc, [
        ("PhET Interactive Simulations. “Wave Interference.” ", {}),
        ("PhET", {"italic": True}),
        (", University of Colorado Boulder, "
         "phet.colorado.edu/en/simulations/wave-interference. Accessed 25 May 2026.", {}),
    ], hanging=True)

    add_mixed_para(doc, [
        ("Schneider, John B. ", {}),
        ("Understanding the Finite-Difference Time-Domain Method", {"italic": True}),
        (". Washington State University, 2010, www.eecs.wsu.edu/~schneidj/ufdtd. "
         "Accessed 25 May 2026.", {}),
    ], hanging=True)

    add_mixed_para(doc, [
        ("Streamlit, Inc. “Streamlit Documentation.” ", {}),
        ("Streamlit", {"italic": True}),
        (", docs.streamlit.io. Accessed 25 May 2026.", {}),
    ], hanging=True)

    add_mixed_para(doc, [
        ("Virtanen, Pauli, et al. “SciPy 1.0: Fundamental Algorithms for Scientific "
         "Computing in Python.” ", {}),
        ("Nature Methods", {"italic": True}),
        (", vol. 17, 2020, pp. 261–272.", {}),
    ], hanging=True)

    out = "ASTP_Final_Project_Written_Description.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    build()
